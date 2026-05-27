"""
Resume Parser Module
Handles PDF, TXT, and DOCX resume parsing
Converts resume files into structured text
"""

import re
import logging
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)


class ResumeParser:
    """Main parser for different resume file formats"""

    SUPPORTED_FORMATS = {".pdf", ".txt", ".docx"}

    @staticmethod
    def parse_resume(filepath: Path) -> Optional[str]:
        """
        Parse resume from file and return extracted text
        Supports PDF, TXT, and DOCX formats
        """
        if not filepath.exists():
            logger.error(f"File not found: {filepath}")
            return None

        suffix = filepath.suffix.lower()

        try:
            if suffix == ".pdf":
                return PDFParser.parse(filepath)
            elif suffix == ".txt":
                return TXTParser.parse(filepath)
            elif suffix == ".docx":
                return DOCXParser.parse(filepath)
            else:
                logger.warning(f"Unsupported file format: {suffix}")
                return None
        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            return None

    @staticmethod
    def get_file_format(filepath: Path) -> Optional[str]:
        """Get file format"""
        suffix = filepath.suffix.lower()
        return suffix if suffix in ResumeParser.SUPPORTED_FORMATS else None


class PDFParser:
    """PDF Resume Parser"""

    @staticmethod
    def parse(filepath: Path) -> Optional[str]:
        """
        Parse PDF resume
        Uses pdfplumber library for extraction
        """
        try:
            import pdfplumber
            
            text_content = []
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        logger.info(f"Extracted page {page_num + 1} from PDF")
            
            return "\n".join(text_content) if text_content else None
            
        except ImportError:
            logger.error("pdfplumber not installed. Using PyPDF2 fallback.")
            return PDFParser._parse_with_pypdf2(filepath)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            return None

    @staticmethod
    def _parse_with_pypdf2(filepath: Path) -> Optional[str]:
        """Fallback PDF parser using PyPDF2"""
        try:
            from PyPDF2 import PdfReader
            
            text_content = []
            with open(filepath, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                        logger.info(f"Extracted page {page_num + 1} with PyPDF2")
            
            return "\n".join(text_content) if text_content else None
            
        except ImportError:
            logger.error("Neither pdfplumber nor PyPDF2 available")
            return None
        except Exception as e:
            logger.error(f"Error parsing PDF with PyPDF2: {e}")
            return None


class TXTParser:
    """TXT Resume Parser"""

    @staticmethod
    def parse(filepath: Path) -> Optional[str]:
        """Parse TXT resume file"""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            logger.info(f"Successfully parsed TXT file: {filepath}")
            return content
        except UnicodeDecodeError:
            try:
                with open(filepath, 'r', encoding='latin-1') as file:
                    content = file.read()
                logger.info(f"Parsed TXT file with latin-1 encoding: {filepath}")
                return content
            except Exception as e:
                logger.error(f"Error parsing TXT with fallback encoding: {e}")
                return None
        except Exception as e:
            logger.error(f"Error parsing TXT file: {e}")
            return None


class DOCXParser:
    """DOCX Resume Parser"""

    @staticmethod
    def parse(filepath: Path) -> Optional[str]:
        """Parse DOCX resume file"""
        try:
            from docx import Document
            
            doc = Document(filepath)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            logger.info(f"Successfully parsed DOCX file: {filepath}")
            return "\n".join(text_content) if text_content else None
            
        except ImportError:
            logger.error("python-docx not installed. Cannot parse DOCX files.")
            return None
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            return None


class ResumeStructurer:
    """Structures parsed resume text into sections"""

    SECTION_PATTERNS = {
        "contact": r"(?:contact|email|phone|address|location|phone\s*number)",
        "summary": r"(?:summary|objective|profile|overview|introduction|about)",
        "experience": r"(?:experience|work\s*experience|employment|professional\s*experience|experience\s*history)",
        "education": r"(?:education|academic|qualification|degree|university|college)",
        "skills": r"(?:skill|technical\s*skill|core\s*competencies|expertise|abilities|proficiency)",
        "certification": r"(?:certification|certificate|certified|award|license)",
        "project": r"(?:project|portfolio|work\s*sample|achievement)",
        "language": r"(?:language|linguistic|multilingual)",
    }

    @staticmethod
    def extract_sections(text: str) -> dict:
        """
        Extract structured sections from resume text
        Returns dict with section names as keys and content as values
        """
        sections = {section: [] for section in ResumeStructurer.SECTION_PATTERNS.keys()}
        sections["other"] = []

        # Split resume into potential section blocks
        lines = text.split('\n')
        current_section = "other"
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            
            # Check if this line is a section header
            for section_name, pattern in ResumeStructurer.SECTION_PATTERNS.items():
                if re.search(pattern, line_clean, re.IGNORECASE) and len(line_clean) < 100:
                    current_section = section_name
                    break
            
            # Add line to current section
            if current_section not in sections:
                sections[current_section] = []
            sections[current_section].append(line_clean)

        return sections

    @staticmethod
    def get_section_text(sections: dict, section_name: str) -> str:
        """Get formatted text for a specific section"""
        if section_name not in sections:
            return ""
        return "\n".join(sections[section_name])


class TextCleaner:
    """Cleans and normalizes extracted text"""

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text
        - Remove extra whitespace
        - Remove special characters
        - Normalize line breaks
        """
        if not text:
            return ""

        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common artifacts
        text = re.sub(r'[^\w\s\-\.\,\@\#\+]', '', text)
        
        return text.strip()

    @staticmethod
    def normalize_spacing(text: str) -> str:
        """Normalize spacing in text"""
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        # Remove spaces before punctuation
        text = re.sub(r' +([.,;:])', r'\1', text)
        return text

    @staticmethod
    def extract_lines(text: str, min_length: int = 2) -> list:
        """
        Extract meaningful lines from text
        Filter out very short lines
        """
        lines = text.split('\n')
        return [line.strip() for line in lines if len(line.strip()) >= min_length]
